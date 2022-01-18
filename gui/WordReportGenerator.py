# -*- coding: utf-8 -*-

# this script aims to generate a testing report in .doxc file

import sys
sys.path.append('gui')
from docx import Document
from docx.shared import Pt
from docx.shared import Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


def generate_word(save_path):
    # Create a word file
    doc = Document()

    # Setting the Margins of the word
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)

    # Inserting an image header to the word
    header_logo = doc.sections[0].header
    paragraph_logo = header_logo.paragraphs[0]
    logo_run = paragraph_logo.add_run()
    logo_run.add_picture("./icon_materials/WordHeaderLogo.png", width=Cm(17.77))

    # Since the python package "python-docx" doesn't support the inserting of the page number, so there is no page number here.

    #Page 1
    #Inserting the topic "Gestrahlte Emission"
    Head = doc.add_heading()
    run1 = Head.add_run("Gestrahlte Emission")
    run1.font.name=u'Times New Roman'
    run1.font.size = Pt(24)
    run1.bold = True
    Head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_break_1 = doc.add_paragraph()

    paragraph1 = doc.add_paragraph()
    run2 = paragraph1.add_run('Allgemeine Informationen')
    run2.font.name = u'Time New Roman'
    run2.font.size = Pt(18)
    run2.bold = True
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = doc.add_table(rows=6, cols=2)
    table1.style ='Light Grid Accent 1'
    table1.rows[0].height=Cm(0.6)
    table1.rows[1].height=Cm(0.6)
    table1.rows[2].height=Cm(0.6)
    table1.rows[3].height=Cm(0.6)
    table1.rows[4].height=Cm(0.6)
    table1.rows[5].height=Cm(0.6)
    table1.cell(0,0).width=Cm(3.76)
    table1.cell(0,1).width=Cm(13.77)
    table1.cell(1,0).width=Cm(3.76)
    table1.cell(1,1).width=Cm(13.77)
    table1.cell(2,0).width=Cm(3.76)
    table1.cell(2,1).width=Cm(13.77)
    table1.cell(3,0).width=Cm(3.76)
    table1.cell(3,1).width=Cm(13.77)
    table1.cell(4,0).width=Cm(3.76)
    table1.cell(4,1).width=Cm(13.77)
    table1.cell(5,0).width=Cm(3.76)
    table1.cell(5,1).width=Cm(13.77)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    run11=table1.cell(0,0).paragraphs[0].add_run('Datendatei')
    run11.font.name='Times New Roman'
    run11.font.size=Pt(13)
    run11.font.color.rgb = RGBColor(66,116,176)

    run12=table1.cell(1,0).paragraphs[0].add_run('Einheit')
    run12.font.name='Times New Roman'
    run12.font.size=Pt(13)
    run12.font.color.rgb = RGBColor(66,116,176)

    run13=table1.cell(2,0).paragraphs[0].add_run('Entfernung')
    run13.font.name='Times New Roman'
    run13.font.size=Pt(13)
    run13.font.color.rgb = RGBColor(66,116,176)

    run14=table1.cell(3,0).paragraphs[0].add_run('Abweichungen')
    run14.font.name='Times New Roman'
    run14.font.size=Pt(13)
    run14.font.color.rgb = RGBColor(66,116,176)

    run15=table1.cell(4,0).paragraphs[0].add_run('Ingenieur*in')
    run15.font.name='Times New Roman'
    run15.font.size=Pt(13)
    run15.font.color.rgb = RGBColor(66,116,176)

    run16=table1.cell(5,0).paragraphs[0].add_run('Zustand')
    run16.font.name='Times New Roman'
    run16.font.size=Pt(13)
    run16.font.color.rgb = RGBColor(66,116,176)

    paragraph_break_2 = doc.add_paragraph()

    paragraph2 = doc.add_paragraph()
    run3 = paragraph2.add_run('EUT und Testbedingungen')
    run3.font.name = u'Times New Roman'
    run3.font.size = Pt(18)
    run3.bold =True
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table2 = doc.add_table(rows=5, cols=2)
    table2.style ='Light Grid Accent 1'
    table2.rows[0].height=Cm(0.6)
    table2.rows[1].height=Cm(0.6)
    table2.rows[2].height=Cm(0.6)
    table2.rows[3].height=Cm(0.6)
    table2.rows[4].height=Cm(0.6)
    table2.cell(0,0).width=Cm(4.05)
    table2.cell(0,1).width=Cm(13.48)
    table2.cell(1,0).width=Cm(4.05)
    table2.cell(1,1).width=Cm(13.48)
    table2.cell(2,0).width=Cm(4.05)
    table2.cell(2,1).width=Cm(13.48)
    table2.cell(3,0).width=Cm(4.05)
    table2.cell(3,1).width=Cm(13.48)
    table2.cell(4,0).width=Cm(4.05)
    table2.cell(4,1).width=Cm(13.48)

    run21=table2.cell(0,0).paragraphs[0].add_run('Kundenname')
    run21.font.name='Times New Roman'
    run21.font.size=Pt(13)
    run21.font.color.rgb = RGBColor(66,116,176)

    run22=table2.cell(1,0).paragraphs[0].add_run('Modellsnummer')
    run22.font.name='Times New Roman'
    run22.font.size=Pt(13)
    run22.font.color.rgb = RGBColor(66,116,176)

    run23=table2.cell(2,0).paragraphs[0].add_run('Ordnungsnummer')
    run23.font.name='Times New Roman'
    run23.font.size=Pt(13)
    run23.font.color.rgb = RGBColor(66,116,176)

    run24=table2.cell(3,0).paragraphs[0].add_run('Beschreibung')
    run24.font.name='Times New Roman'
    run24.font.size=Pt(13)
    run24.font.color.rgb = RGBColor(66,116,176)

    run25=table2.cell(4,0).paragraphs[0].add_run('Bedingungen')
    run25.font.name='Times New Roman'
    run25.font.size=Pt(13)
    run25.font.color.rgb = RGBColor(66,116,176)

    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    paragraph_break_3 = doc.add_paragraph()

    paragraph3 = doc.add_paragraph()
    run4 = paragraph3.add_run('Benutzernotizen')
    run4.font.name = u'Times New Roman'
    run4.font.size = Pt(18)
    run4.bold = True
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table3 = doc.add_table(rows=1, cols=2)
    table3.style ='Light Grid Accent 1'
    table3.cell(0,0).width=Cm(3.76)
    table3.cell(0,1).width=Cm(13.77)

    run31=table3.cell(0,0).paragraphs[0].add_run('TT.MM.JJJJ\n 00:00:00')
    run31.font.name='Times New Roman'
    run31.font.size=Pt(13)
    run31.font.color.rgb = RGBColor(66,116,176)

    table3.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()

    #Page 2

    paragraph_break_4 = doc.add_paragraph()

    paragraph4 = doc.add_paragraph()
    run5 = paragraph4.add_run('Diagramm')
    run5.font.name = u'Times New Roman'
    run5.font.size = Pt(18)
    run5.bold = True
    paragraph4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(save_path, width=Cm(17.75))

    doc.add_page_break()

    #Page 3
    paragraph_break_5 = doc.add_paragraph()

    paragraph5 = doc.add_paragraph()
    run6 = paragraph5.add_run('Einstellungen des Umgebungsempfängers',)
    run6.font.name = u'Times New Roman'
    run6.font.size = Pt(18)
    run6.bold = True
    paragraph5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph_break_6 = doc.add_paragraph()

    paragraph6 = doc.add_paragraph()
    run7 = paragraph6.add_run('Date Range:     DD.MM.JJ     00:00:00   --   DD.MM.JJ     00:00:0')
    run7.font.name = u'Times New Roman'
    run7.font.size = Pt(13)
    run7.bold = False

    table4 = doc.add_table(rows=2, cols=6)
    table4.style ='Light Grid Accent 1'
    table4.rows[0].height=Cm(0.8)
    table4.rows[1].height=Cm(0.8)
    table4.cell(0,0).width=Cm(2.92)
    table4.cell(0,1).width=Cm(2.92)
    table4.cell(0,2).width=Cm(2.92)
    table4.cell(0,3).width=Cm(2.92)
    table4.cell(0,4).width=Cm(2.92)
    table4.cell(0,5).width=Cm(2.92)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    run41=table4.cell(0,0).paragraphs[0].add_run('Effektive\nFrequenzen')
    run41.font.name='Times New Roman'
    run41.font.size=Pt(13)
    run41.font.color.rgb = RGBColor(66,116,176)

    run42=table4.cell(0,1).paragraphs[0].add_run('6dB Bandbreite')
    run42.font.name='Times New Roman'
    run42.font.size=Pt(13)
    run42.font.color.rgb = RGBColor(66,116,176)

    run43=table4.cell(0,2).paragraphs[0].add_run('Verweilzeit')
    run43.font.name='Times New Roman'
    run43.font.size=Pt(13)
    run43.font.color.rgb = RGBColor(66,116,176)

    run44=table4.cell(0,3).paragraphs[0].add_run('Referenzniveau')
    run44.font.name='Times New Roman'
    run44.font.size=Pt(13)
    run44.font.color.rgb = RGBColor(66,116,176)

    run45=table4.cell(0,4).paragraphs[0].add_run('Detektor')
    run45.font.name='Times New Roman'
    run45.font.size=Pt(13)
    run45.font.color.rgb = RGBColor(66,116,176)

    run46=table4.cell(0,5).paragraphs[0].add_run('PreAmp-Zustand')
    run46.font.name='Times New Roman'
    run46.font.size=Pt(13)
    run46.font.color.rgb = RGBColor(66,116,176)

    table4.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,2).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,3).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,4).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,5).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table4.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table4.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table4.cell(0,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table4.cell(0,3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table4.cell(0,4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table4.cell(0,5).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    paragraph_break_7 = doc.add_paragraph()

    paragraph7 = doc.add_paragraph()
    run8 = paragraph7.add_run('Verwendete Umgebungsausrüstungen')
    run8.font.name = u'Times New Roman'
    run8.font.size = Pt(18)
    run8.bold = True
    paragraph7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table5 = doc.add_table(rows=13, cols=2)
    table5.style ='Light Grid Accent 1'
    table5.rows[0].height=Cm(0.5)
    table5.rows[1].height=Cm(0.5)
    table5.rows[2].height=Cm(0.5)
    table5.rows[3].height=Cm(0.5)
    table5.rows[4].height=Cm(0.5)
    table5.rows[5].height=Cm(0.5)
    table5.rows[6].height=Cm(0.5)
    table5.rows[7].height=Cm(0.5)
    table5.rows[8].height=Cm(0.5)
    table5.rows[9].height=Cm(0.5)
    table5.rows[10].height=Cm(0.5)
    table5.rows[11].height=Cm(0.5)
    table5.rows[12].height=Cm(0.5)
    table5.cell(0,0).width=Cm(5.13)
    table5.cell(0,1).width=Cm(12.4)
    table5.cell(1,0).width=Cm(5.13)
    table5.cell(1,1).width=Cm(12.4)
    table5.cell(2,0).width=Cm(5.13)
    table5.cell(2,1).width=Cm(12.4)
    table5.cell(3,0).width=Cm(5.13)
    table5.cell(3,1).width=Cm(12.4)
    table5.cell(4,0).width=Cm(5.13)
    table5.cell(4,1).width=Cm(12.4)
    table5.cell(5,0).width=Cm(5.13)
    table5.cell(5,1).width=Cm(12.4)
    table5.cell(6,0).width=Cm(5.13)
    table5.cell(6,1).width=Cm(12.4)
    table5.cell(7,0).width=Cm(5.13)
    table5.cell(7,1).width=Cm(12.4)
    table5.cell(8,0).width=Cm(5.13)
    table5.cell(8,1).width=Cm(12.4)
    table5.cell(9,0).width=Cm(5.13)
    table5.cell(9,1).width=Cm(12.4)
    table5.cell(10,0).width=Cm(5.13)
    table5.cell(10,1).width=Cm(12.4)
    table5.cell(11,0).width=Cm(5.13)
    table5.cell(11,1).width=Cm(12.4)
    table5.cell(12,0).width = Cm(5.13)
    table5.cell(12,1).width=Cm(12.4)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER

    run51=table5.cell(0,0).paragraphs[0].add_run('Effektive Frequenzen')
    run51.font.name='Times New Roman'
    run51.font.size=Pt(13)
    run51.font.color.rgb = RGBColor(66,116,176)

    run52=table5.cell(1,0).paragraphs[0].add_run('Effektive Daten')
    run52.font.name='Times New Roman'
    run52.font.size=Pt(13)
    run52.font.color.rgb = RGBColor(66,116,176)

    run53=table5.cell(2,0).paragraphs[0].add_run('Empfänger')
    run53.font.name='Times New Roman'
    run53.font.size=Pt(13)
    run53.font.color.rgb = RGBColor(66,116,176)

    run54=table5.cell(3,0).paragraphs[0].add_run('Externer Vorverstärker')
    run54.font.name='Times New Roman'
    run54.font.size=Pt(13)
    run54.font.color.rgb = RGBColor(66,116,176)

    run55=table5.cell(4,0).paragraphs[0].add_run('Antenne')
    run55.font.name='Times New Roman'
    run55.font.size=Pt(13)
    run55.font.color.rgb = RGBColor(66,116,176)

    run56=table5.cell(5,0).paragraphs[0].add_run('Antennenkontroller')
    run56.font.name='Times New Roman'
    run56.font.size=Pt(13)
    run56.font.color.rgb = RGBColor(66,116,176)

    run57=table5.cell(6,0).paragraphs[0].add_run('Drehtabelle')
    run57.font.name='Times New Roman'
    run57.font.size=Pt(13)
    run57.font.color.rgb = RGBColor(66,116,176)

    run58=table5.cell(7,0).paragraphs[0].add_run('RF-Kabel 1')
    run58.font.name='Times New Roman'
    run58.font.size=Pt(13)
    run58.font.color.rgb = RGBColor(66,116,176)

    run59=table5.cell(8,0).paragraphs[0].add_run('RF-Kabel 2')
    run59.font.name='Times New Roman'
    run59.font.size=Pt(13)
    run59.font.color.rgb = RGBColor(66,116,176)

    run591=table5.cell(9,0).paragraphs[0].add_run('Abschwächer 1')
    run591.font.name='Times New Roman'
    run591.font.size=Pt(13)
    run591.font.color.rgb = RGBColor(66,116,176)

    run592=table5.cell(10,0).paragraphs[0].add_run('Abschwächer 2')
    run592.font.name='Times New Roman'
    run592.font.size=Pt(13)
    run592.font.color.rgb = RGBColor(66,116,176)

    run593=table5.cell(11,0).paragraphs[0].add_run('Systemkontroller 1')
    run593.font.name='Times New Roman'
    run593.font.size=Pt(13)
    run593.font.color.rgb = RGBColor(66,116,176)

    run594=table5.cell(12,0).paragraphs[0].add_run('Systemkontroller 2')
    run594.font.name='Times New Roman'
    run594.font.size=Pt(13)
    run594.font.color.rgb = RGBColor(66,116,176)


    #Page4
    doc.add_page_break()
    paragraph_break_8 = doc.add_paragraph()

    doc.save('D:/ss.docx')

if __name__ == "__main__":
    generate_word(save_path='D:/demo.png')
